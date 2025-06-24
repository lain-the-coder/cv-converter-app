# cv_converter.py - Enhanced CV Converter with Authentication
# -----------------------------------------------------------------
# pip install streamlit PyPDF2 python-docx google-generativeai
import os, re, json, time
from io import BytesIO
from typing import Dict, Any, List
from datetime import datetime
import streamlit as st
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai

# ────────────────────────────────────────────────────────────────
#  Page Configuration
# ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="CV Converter", 
    page_icon="📄", 
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ────────────────────────────────────────────────────────────────
#  Configuration
# ────────────────────────────────────────────────────────────────
DEFAULT_API_KEY = ""  # Remove hardcoded key for production
# ────────────────────────────────────────────────────────────────
#  Authentication Functions
# ────────────────────────────────────────────────────────────────
def check_company_email():
    """Verify user has company email domain and password."""
    
    # Get company domain and password from secrets
    try:
        COMPANY_DOMAIN = st.secrets["company_domain"]  # e.g., "@yourcompany.com"
        APP_PASSWORD = st.secrets["app_password"]  # Common password for the team
    except:
        st.error("⚠️ Company domain or password not configured. Contact administrator.")
        st.stop()
    
    # Initialize authentication state
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    
    # If not authenticated, show login form
    if not st.session_state.authenticated:
        st.markdown("## 🔐 CV Converter Login Page")
        st.markdown("Please authenticate with your company email and password to access the CV converter.")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col2:
            email = st.text_input("Enter your company email address:", 
                                placeholder=f"yourname{COMPANY_DOMAIN}")
            password = st.text_input("Enter password:", type="password")
            
            if st.button("Access CV Converter", type="primary", use_container_width=True):
                if email.lower().endswith(COMPANY_DOMAIN.lower()) and password == APP_PASSWORD:
                    st.session_state.authenticated = True
                    st.session_state.user_email = email
                    st.session_state.login_time = datetime.now()
                    
                    # Log successful access
                    log_access(email, "login_success")
                    
                    st.success(f"✅ Welcome {email}!")
                    st.rerun()
                else:
                    # Log failed access attempt
                    log_access(email, "login_failed")
                    if not email.lower().endswith(COMPANY_DOMAIN.lower()):
                        st.error(f"❌ Access restricted to {COMPANY_DOMAIN} emails only")
                    else:
                        st.error("❌ Invalid password")
                    
        st.markdown("---")
        st.caption("This tool is for internal HR use only. Unauthorized access is prohibited.")
        
    return st.session_state.authenticated

def log_access(email: str, action: str, details: str = ""):
    """Log user actions for audit trail."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # Create log entry
    log_entry = {
        "timestamp": timestamp,
        "email": email,
        "action": action,
        "details": details
    }
    
    # In production, you might want to send this to a logging service
    # For now, we'll just print to console (visible in Streamlit logs)
    print(f"[ACCESS LOG] {timestamp} - {email} - {action} - {details}")

def check_session_timeout():
    """Check if session has timed out (30 minutes)."""
    if "login_time" in st.session_state:
        elapsed = datetime.now() - st.session_state.login_time
        if elapsed.total_seconds() > 1800:  # 30 minutes
            st.warning("⏱️ Session expired. Please login again.")
            for key in ["authenticated", "user_email", "login_time"]:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()

# ────────────────────────────────────────────────────────────────
#  Original Helper Functions
# ────────────────────────────────────────────────────────────────
def mask_api_key(api_key: str) -> str:
    """Mask API key for display, showing only first 4 and last 4 characters."""
    if not api_key or len(api_key) < 12:
        return api_key
    return f"{api_key[:4]}{'*' * (len(api_key) - 8)}{api_key[-4:]}"

def extract_text(upload) -> str:
    try:
        if upload.type == "application/pdf":
            return "\n".join(p.extract_text() or "" for p in PyPDF2.PdfReader(upload).pages)
        if upload.type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return "\n".join(p.text for p in Document(BytesIO(upload.getvalue())).paragraphs)
        return upload.read().decode("utf-8", errors="ignore")
    except Exception as e:
        st.error(f"Error reading {upload.name}: {e}")
        return ""

def format_date(date_str: str) -> str:
    """Convert various date formats to MMM YYYY format."""
    if not date_str:
        return ""
    
    # Check for present/current/ongoing terms
    if date_str.lower() in ["present", "current", "ongoing", "till date", "now", "till now"]:
        return "Present"
    
    # Common date patterns
    patterns = [
        # Handle Sep-2015, Sep 2015, Sep/2015 formats
        (r'(\w{3,})[- /](\d{4})', lambda m: f"{m.group(1).upper()[:3]} {m.group(2)}"),
        # Handle 09/2015, 09-2015 formats  
        (r'(\d{1,2})[/-](\d{4})', lambda m: f"{get_month_abbr(m.group(1).zfill(2))} {m.group(2)}"),
        # Handle September 2015, Sep 2015 formats
        (r'(\w+)\s+(\d{4})', lambda m: f"{m.group(1)[:3].upper()} {m.group(2)}"),
        # Handle September, 2015 formats
        (r'(\w+),?\s+(\d{4})', lambda m: f"{m.group(1)[:3].upper()} {m.group(2)}"),
    ]
    
    for pattern, formatter in patterns:
        match = re.search(pattern, date_str, re.IGNORECASE)
        if match:
            return formatter(match)
    
    return date_str  # Return as-is if no pattern matches

def get_month_abbr(month_num: str) -> str:
    """Convert month number to 3-letter abbreviation."""
    months = ["JAN", "FEB", "MAR", "APR", "MAY", "JUN", 
              "JUL", "AUG", "SEP", "OCT", "NOV", "DEC"]
    try:
        # Handle both "02" and "2" formats
        month_int = int(month_num.lstrip('0') or '0')
        if 1 <= month_int <= 12:
            return months[month_int - 1]
    except:
        pass
    return month_num

def format_duration(duration: str, is_first_experience: bool = False) -> str:
    """Format duration string to MMM YYYY - MMM YYYY format."""
    if not duration:
        return ""
    
    # Preserve "Present" for ongoing positions
    if " - Present" in duration or "- Present" in duration:
        # Just format the start date part
        parts = re.split(r'\s*-\s*', duration)
        if parts:
            start = parts[0].strip()
            # Convert month to uppercase
            start = re.sub(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', 
                          lambda m: m.group(1).upper(), start, flags=re.IGNORECASE)
            # Ensure space between month and year
            start = re.sub(r'([A-Z]{3})-?(\d{4})', r'\1 \2', start)
            return f"{start} - Present"
    
    # Split by common separators
    parts = re.split(r'\s*[-–—]\s*', duration)
    
    if len(parts) == 2:
        start = format_date(parts[0].strip())
        end = format_date(parts[1].strip())
        return f"{start} - {end}"
    elif len(parts) == 1 and is_first_experience:
        # For first experience, if only start date, add Present
        start = format_date(parts[0].strip())
        if start and start != "Present":
            return f"{start} - Present"
        return start
    
    return duration

def format_name(name: str) -> str:
    """Convert name from ALL CAPS to Proper Case."""
    if not name:
        return ""
    
    # Handle common name patterns
    words = name.split()
    formatted_words = []
    
    for word in words:
        if word.isupper() and len(word) > 1:
            # Convert from ALL CAPS to Proper Case
            formatted_words.append(word.capitalize())
        else:
            formatted_words.append(word)
    
    return " ".join(formatted_words)

# ────────────────────────────────────────────────────────────────
#  Enhanced Gemini wrapper for comprehensive extraction
# ────────────────────────────────────────────────────────────────
class CVExtractor:
    def __init__(self, api_key: str):
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel("models/gemini-1.5-flash-latest")
        self.cfg = {"temperature": 0.1, "top_p": 0.1, "top_k": 1}

    def extract(self, cv_text: str) -> Dict[str, Any]:
        prompt = f"""Extract comprehensive information from this CV and return as JSON.


CRITICAL INSTRUCTIONS:
1. Extract ALL experiences (up to 20, most recent first)
2. IMPORTANT: If a person has multiple projects at the SAME COMPANY, combine them into ONE experience entry:
   - Company name should be just the company name (e.g., "Seertree Global Services")
   - Duration should span from the earliest project start to the latest project end
   - In responsibilities section, list each project as a subheading followed by its responsibilities
   - Format: "Project Name: [Name], Location: [Location], Duration: [Project Start - Project End]" as the first line, then responsibilities below
   - Look for project-specific dates in the CV (these may be different from overall company duration)
3. For each experience, you MUST include:
   - Company name: Always include location if mentioned anywhere in the experience
     * Format: "Company Name, Location: Location"
     * Look for location in company line, job title line, or anywhere in the experience section
     * Examples: "ITForce Technology, Location: Dubai", "Vogue International FZE, Location: Sharjah"
     * Location can be city, country, or both (e.g., "Dubai", "Sharjah", "Dubai, UAE", "Delhi, India")
   - Job title/role (in proper case, not ALL CAPS)
   - Duration: Use the date format provided
   - Responsibilities: 
     * If multiple projects at same company, structure with project headers including individual project durations
     * Otherwise, list responsibilities directly
     * Extract EVERY responsibility mentioned
     * If there's an "Environment:" or "Technologies:" or "Versions:" section, add it as the LAST bullet point
     * Format environment info as: "Environment/Technologies: [list all tools, versions, technologies mentioned]"
     * Do NOT add bullet points - they will be added by the template
4. Extract ALL technical skills comprehensively
5. Extract ALL certifications with their FULL names and IDs if mentioned
6. If no professional summary exists, create one based on the CV content
7. For candidate name, use proper case (e.g., "Raju Gujar" not "RAJU GUJAR")
8. For position/role, use proper case (e.g., "Senior Technical Engineer" not "SENIOR TECHNICAL ENGINEER")

Return this exact JSON structure:
{{
  "candidate_name": "Full name in proper case",
  "position": "Current or most recent job title in proper case",
  "education": "Highest degree with field and university",
  "total_experience_years": "Number only (e.g., 11)",
  "phone": "Phone number with country code if present",
  "email": "Email address",
  "intro_paragraph": "Professional summary add as many sentences as available in CV in paragraph format, word it in a structured manner, summarize if needed.",
  "experiences": [
    {{
      "company": "Company name only (no project details here)",
      "role": "Job title in proper case",
      "duration": "Start date of first project - End date of last project (or Present)",
      "responsibilities": [
        "Project Name: Project 1 Name, Location: Location 1, Duration: Sep 2015 - Jun 2018",
        "First responsibility for project 1",
        "Second responsibility for project 1",
        "All other responsibilities for project 1",
        "Environment: Oracle 19c/12c, SQL * Plus, TOAD, SQL*Loader, SQL Developer, Shell Scripts, UNIX, Windows 10",
        "Project Name: Project 2 Name, Location: Location 2, Duration: Jul 2018 - Present",
        "First responsibility for project 2",
        "Second responsibility for project 2",
        "Environment: Oracle 19c/12c, SQL * Plus, TOAD, SQL*Loader, SQL Developer, Shell Scripts, UNIX, Windows 10",
        "Continue for all projects and responsibilities"
      ]
    }}
  ],
  "technical_skills": ["List ALL technical skills mentioned"],
  "certifications": ["Full certification names with IDs"],
  "language_skills": ["Language - Proficiency level"]
}}

EXAMPLE for someone with multiple projects at same company:
{{
  "experiences": [
    {{
      "company": "Seertree Global Services",
      "role": "Technical Consultant",
      "duration": "Sep-2015 - Present",
      "responsibilities": [
        "Project Name: GE Appliances, Location: Offshore, Duration: Nov'23 – Till Now",
        "Developed the custom packages to create credit card and assigning customer",
        "Modified the standard package based on the client requirement",
        "Using form personalization showing the credit card details at Sales order form",
        "Developed custom reports as per customer requirements",
        "Project Name: Fine Hygienic Holding (FHH), Location: Offshore, Duration: Sep'22 – Nov'23",
        "Migrating data from EBS to fusion using payloads, Rest API's and Spread sheets",
        "Modified standard report data models and layouts in fusion based on client requirements",
        "Developed custom BI reports in fusion",
        "Worked in BI bursting for sending email's and sent output to printer in fusion"
      ]
    }}
  ]
}}

EXAMPLE for someone with single job (no projects):
{{
  "experiences": [
    {{
      "company": "ITForce Technology, Location: Dubai",
      "role": "Senior Technical Engineer",
      "duration": "Feb 2023 - Present",
      "responsibilities": [
        "Administered and optimized Office 365 applications",
        "Enhanced security protocols by managing Barracuda Email Security Gateway",
        "Led data migration projects with a focus on accuracy and efficiency"
      ]
    }}
  ]
}}

EXAMPLE of location extraction from actual CV text:
If CV shows: "Vogue International FZE, Sharjah"
Extract as: "company": "Vogue International FZE, Location: Sharjah"

If CV shows: "HP (Hewlett Packard) payroll of Metalogic PVT, Delhi, India"  
Extract as: "company": "HP (Hewlett Packard) payroll of Metalogic PVT, Location: Delhi, India"

CV TEXT:
{cv_text}

IMPORTANT REMINDERS:
- Consolidate multiple projects at the same company into ONE experience entry
- ALWAYS extract location from the CV and include it in company field as: "Company Name, Location: Location"
- Location appears in various formats - after company name, on separate line, or with country
- Look for cities like Dubai, Sharjah, Delhi, and countries like UAE, India
- If location is split (e.g., "Dubai" on one line, "UAE" on another), combine them as "Dubai, UAE"
- Extract ALL locations mentioned - if multiple cities/countries, include all
- If multiple projects exist at same company:
  * List each project as a subheading in responsibilities: "Project Name: [Name], Location: [Location], Duration: [Start - End]"
  * Extract project-specific durations from the CV when available
  * Calculate total duration from first project start to last project end
  * Do NOT create separate experience entries for different projects at the same company
- Extract ALL responsibilities under their respective companies or project subheadings
- Do NOT add bullets or special characters to responsibilities - plain text only
- Ensure responsibilities are correctly attributed to their respective roles/projects
- If no project names are mentioned, just list responsibilities directly without project headers
- Look for sections labeled "Environment:", "Technologies:", "Tools:", "Versions:", "Tech Stack:" etc.
- These should be captured as the last item in the responsibilities list for that experience
- Preserve the exact format and all items listed
- Common patterns to look for:
  * "Environment: Oracle 19c/12c, SQL * Plus..."
  * "Versions: • Oracle Apps R12.1.1 • Oracle Database 10g..."
  * "Technologies used: Java, Spring Boot, MySQL..."
- For overall company duration: if CV says "Sep-2015 to till date" extract as "Sep-2015 - Present"
- Preserve date formats as they appear but ensure "Present" is used for ongoing positions

RETURN ONLY THE JSON:"""

        try:
            r = self.model.generate_content(prompt, generation_config=self.cfg)
            raw = re.sub(r'```(?:json)?', '', r.text).strip('`')
            
            # Extract JSON
            match = re.search(r'\{.*\}', raw, re.DOTALL)
            if match:
                data = json.loads(match.group(0))
                return self._validate_data(data)
            else:
                raise ValueError("No JSON found")
                
        except Exception as e:
            st.warning(f"⚠️ Extraction error: {str(e)}")
            return self._get_empty_data()        

    def _validate_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Ensure data structure is complete and properly formatted."""
        # Format candidate name and position
        if "candidate_name" in data:
            data["candidate_name"] = format_name(data["candidate_name"])
        
        if "position" in data:
            data["position"] = format_name(data["position"])
        
        # Default language skills if not found
        if "language_skills" not in data or not data["language_skills"]:
            data["language_skills"] = ["English - Fluent"]
        
        # Ensure experiences exist and have proper structure
        if "experiences" not in data:
            data["experiences"] = []
        
        # Format existing experiences
        for i, exp in enumerate(data["experiences"]):
            if "role" in exp:
                exp["role"] = format_name(exp["role"])
            if "duration" in exp:
                # Special handling for first experience
                exp["duration"] = format_duration(exp["duration"], is_first_experience=(i == 0))
        
        # IMPORTANT: Don't pad experiences anymore - keep only actual experiences
        # This allows the fill_template function to handle row deletion
        
        # Ensure each experience has all fields
        for exp in data["experiences"]:
            if "company" not in exp or not exp["company"]:
                exp["company"] = ""
            if "role" not in exp or not exp["role"]:
                exp["role"] = ""
            if "duration" not in exp or not exp["duration"]:
                exp["duration"] = ""
            if "responsibilities" not in exp or not isinstance(exp["responsibilities"], list):
                exp["responsibilities"] = []
        
        return data

    def _get_empty_data(self) -> Dict[str, Any]:
        return {
            "candidate_name": "",
            "position": "",
            "education": "",
            "total_experience_years": "",
            "phone": "",
            "email": "",
            "intro_paragraph": "",
            "experiences": [],  # Empty list, no padding
            "technical_skills": [],
            "certifications": [],
            "language_skills": ["English - Fluent"]
        }

# ────────────────────────────────────────────────────────────────
#  Helper: Check if a table row contains experience placeholders
# ────────────────────────────────────────────────────────────────
def contains_experience_placeholder(text: str, exp_num: int) -> bool:
    """Check if text contains placeholders for a specific experience number."""
    patterns = [
        f"{{{{EXP{exp_num}_COMPANY}}}}",
        f"{{{{EXP{exp_num}_ROLE}}}}",
        f"{{{{EXP{exp_num}_DURATION}}}}",
        f"{{{{EXP{exp_num}_RESP"
    ]
    return any(pattern in text for pattern in patterns)

def get_row_text(row) -> str:
    """Get all text from a table row."""
    text = ""
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            text += paragraph.text + " "
    return text

def should_delete_row(row, exp_num: int, has_data: bool) -> bool:
    """Determine if a row should be deleted based on experience data availability."""
    row_text = get_row_text(row)
    
    # Check if this row contains placeholders for this experience number
    if contains_experience_placeholder(row_text, exp_num):
        # If no data for this experience, mark for deletion
        return not has_data
    
    return False

# ────────────────────────────────────────────────────────────────
#  Enhanced template filling with row deletion
# ────────────────────────────────────────────────────────────────
def set_paragraph_format(paragraph, text, font_name="Arial", font_size=11, bold=False):
    """Set consistent formatting for a paragraph."""
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black

def fill_template(doc: Document, d: Dict[str, Any]) -> Document:
    """Fill template with proper formatting and delete unused experience rows."""
    
    # Basic replacements
    basic_repl = {
        "{{CANDIDATE_NAME}}": d.get("candidate_name", ""),
        "{{POSITION}}": d.get("position", ""),
        "{{EDUCATION}}": d.get("education", ""),
        "{{TOTAL_EXPERIENCE_YEARS}}": str(d.get("total_experience_years", "")),
        "{{PHONE}}": d.get("phone", ""),
        "{{EMAIL}}": d.get("email", ""),
        "{{INTRO_PARAGRAPH}}": d.get("intro_paragraph", ""),
    }
    
    # Track which experiences have data
    experiences_with_data = set()
    
    # Process experiences (up to 20)
    exp_repl = {}
    for i in range(1, 21):  # 1 to 20
        if i <= len(d.get("experiences", [])):
            exp = d["experiences"][i-1]
            # Check if this experience has meaningful data
            if exp.get("company") and exp.get("role"):
                experiences_with_data.add(i)
                
                # Mark company placeholders for bold formatting
                company_text = exp.get("company", "")
                exp_repl[f"{{{{EXP{i}_COMPANY}}}}"] = f"<<<BOLD>>>{company_text}<<<END_BOLD>>>"
                exp_repl[f"{{{{EXP{i}_ROLE}}}}"] = exp.get("role", "")
                exp_repl[f"{{{{EXP{i}_DURATION}}}}"] = exp.get("duration", "")
                
                # Handle responsibilities
                responsibilities = exp.get("responsibilities", [])
                
                # Handle all 100 responsibility placeholders
                for j in range(1, 101):  # 1 to 100
                    placeholder = f"{{{{EXP{i}_RESP{j}}}}}"
                    if j <= len(responsibilities):
                        exp_repl[placeholder] = responsibilities[j-1]
                    else:
                        # Mark empty responsibilities for removal
                        exp_repl[placeholder] = "<<<REMOVE_THIS_LINE>>>"
            else:
                # Mark all placeholders for this experience for deletion
                exp_repl[f"{{{{EXP{i}_COMPANY}}}}"] = "<<<DELETE_EXPERIENCE>>>"
                exp_repl[f"{{{{EXP{i}_ROLE}}}}"] = "<<<DELETE_EXPERIENCE>>>"
                exp_repl[f"{{{{EXP{i}_DURATION}}}}"] = "<<<DELETE_EXPERIENCE>>>"
                for j in range(1, 101):
                    exp_repl[f"{{{{EXP{i}_RESP{j}}}}}"] = "<<<DELETE_EXPERIENCE>>>"
        else:
            # No data for this experience - mark for deletion
            exp_repl[f"{{{{EXP{i}_COMPANY}}}}"] = "<<<DELETE_EXPERIENCE>>>"
            exp_repl[f"{{{{EXP{i}_ROLE}}}}"] = "<<<DELETE_EXPERIENCE>>>"
            exp_repl[f"{{{{EXP{i}_DURATION}}}}"] = "<<<DELETE_EXPERIENCE>>>"
            for j in range(1, 101):
                exp_repl[f"{{{{EXP{i}_RESP{j}}}}}"] = "<<<DELETE_EXPERIENCE>>>"
    
    # Format skills and certifications as bullet points
    tech_skills = d.get("technical_skills", [])
    if tech_skills:
        tech_skills_text = "\n".join([f"• {skill}" for skill in tech_skills])
    else:
        tech_skills_text = ""
    
    certs = d.get("certifications", [])
    if certs:
        certs_text = "\n".join([f"• {cert}" for cert in certs])
    else:
        certs_text = "N/A"  # No bullet for N/A
    
    langs = d.get("language_skills", [])
    if langs:
        langs_text = ", ".join(langs)
    else:
        langs_text = "English - Fluent"
    
    # Add formatted lists to replacements
    basic_repl["{{TECHNICAL_SKILLS_LIST}}"] = tech_skills_text
    basic_repl["{{CERTIFICATIONS_LIST}}"] = certs_text
    basic_repl["{{LANGUAGE_SKILLS_LIST}}"] = langs_text
    
    # Combine all replacements
    all_repl = {**basic_repl, **exp_repl}
    
    # Process paragraphs
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        new_text = original_text
        
        # Apply replacements
        for placeholder, value in all_repl.items():
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, value)
        
        # Check if this paragraph is part of a deleted experience section
        if "<<<DELETE_EXPERIENCE>>>" in new_text:
            paragraphs_to_remove.append(paragraph)
            continue
        
        # Check if this line should be removed (empty responsibility)
        if "<<<REMOVE_THIS_LINE>>>" in new_text:
            # For regular paragraphs, only remove if it's just the marker
            if new_text.strip() in ["<<<REMOVE_THIS_LINE>>>", "- <<<REMOVE_THIS_LINE>>>"]:
                paragraphs_to_remove.append(paragraph)
                continue
            else:
                # Replace the marker with empty string in the text
                new_text = new_text.replace("<<<REMOVE_THIS_LINE>>>", "")
        
        # If text changed, update with formatting
        if new_text != original_text:
            paragraph.clear()
            # Remove the bullet point if the line only contains "-"
            if new_text.strip() != "-":
                # Check for bold markers
                if "<<<BOLD>>>" in new_text and "<<<END_BOLD>>>" in new_text:
                    # Extract and apply bold formatting
                    parts = new_text.split("<<<BOLD>>>")
                    for i, part in enumerate(parts):
                        if i == 0 and part:
                            # Text before first bold marker
                            run = paragraph.add_run(part)
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        elif "<<<END_BOLD>>>" in part:
                            # This part contains bold text
                            bold_parts = part.split("<<<END_BOLD>>>")
                            # Add bold text
                            run = paragraph.add_run(bold_parts[0])
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            # Add remaining text if any
                            if len(bold_parts) > 1 and bold_parts[1]:
                                run = paragraph.add_run(bold_parts[1])
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # No bold markers, normal text
                    paragraph.add_run(new_text)
                    # Apply Arial 11pt black formatting
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0)
                # Set 1.5 line spacing
                paragraph.paragraph_format.line_spacing = 1.5
    
    # Remove empty paragraphs
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)
    
    # Process tables and handle row deletion
    for table in doc.tables:
        rows_to_delete = []
        
        # First pass: identify rows to delete
        for row_idx, row in enumerate(table.rows):
            row_text = get_row_text(row)
            
            # Check each experience number
            for exp_num in range(1, 21):
                if contains_experience_placeholder(row_text, exp_num):
                    if exp_num not in experiences_with_data:
                        # This row contains placeholders for an experience we don't have data for
                        rows_to_delete.append(row_idx)
                        break
        
        # Second pass: process cells in rows we're keeping
        for row_idx, row in enumerate(table.rows):
            if row_idx in rows_to_delete:
                continue  # Skip rows marked for deletion
                
            for cell in row.cells:
                paragraphs_to_remove = []
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    new_text = original_text
                    
                    # Apply replacements
                    for placeholder, value in all_repl.items():
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, value)
                    
                    # Skip if it's marked for deletion
                    if "<<<DELETE_EXPERIENCE>>>" in new_text:
                        continue
                    
                    # Check if this line should be removed
                    if "<<<REMOVE_THIS_LINE>>>" in new_text:
                        # Check if it's a complete line with just the marker and possibly a bullet
                        if new_text.strip() in ["<<<REMOVE_THIS_LINE>>>", "- <<<REMOVE_THIS_LINE>>>"]:
                            paragraphs_to_remove.append(paragraph)
                            continue
                        else:
                            # Replace the marker with empty string
                            new_text = new_text.replace("<<<REMOVE_THIS_LINE>>>", "")
                    
                    # If text changed, update with formatting
                    if new_text != original_text:
                        paragraph.clear()
                        
                        # Skip if it's just a bullet point with no content
                        if new_text.strip() == "-":
                            paragraphs_to_remove.append(paragraph)
                            continue
                        
                        # Check for bold markers in table cells
                        if "<<<BOLD>>>" in new_text and "<<<END_BOLD>>>" in new_text:
                            # Extract and apply bold formatting
                            parts = new_text.split("<<<BOLD>>>")
                            for i, part in enumerate(parts):
                                if i == 0 and part:
                                    # Text before first bold marker
                                    run = paragraph.add_run(part)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                elif "<<<END_BOLD>>>" in part:
                                    # This part contains bold text
                                    bold_parts = part.split("<<<END_BOLD>>>")
                                    # Add bold text
                                    run = paragraph.add_run(bold_parts[0])
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                    # Add remaining text if any
                                    if len(bold_parts) > 1 and bold_parts[1]:
                                        run = paragraph.add_run(bold_parts[1])
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(11)
                                        run.font.color.rgb = RGBColor(0, 0, 0)
                        # Handle multi-line content (like skills/certs)
                        elif '\n' in new_text:
                            lines = new_text.split('\n')
                            for idx, line in enumerate(lines):
                                if idx > 0:
                                    paragraph = cell.add_paragraph()
                                
                                # Check if this is a project header line with Duration
                                if (line.strip().startswith("Project Name:") and 
                                    ("Location:" in line.strip() and "Duration:" in line.strip())):
                                    # Make the entire project header bold (including Duration)
                                    run = paragraph.add_run(line)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.bold = True  # Make entire line bold
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                elif line.strip().startswith("Project Name:") and "Location:" in line.strip():
                                    # Old format without Duration - still make bold
                                    run = paragraph.add_run(line)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                else:
                                    run = paragraph.add_run(line)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.bold = False
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                # Set 1.5 line spacing
                                paragraph.paragraph_format.line_spacing = 1.5
                        else:
                            # Single line handling
                            if (new_text.strip().startswith("Project Name:") and 
                                ("Location:" in new_text.strip() and "Duration:" in new_text.strip())):
                                # Project header with Duration - make entire line bold
                                if new_text.strip().startswith("• "):
                                    new_text = new_text.replace("• ", "", 1)
                                run = paragraph.add_run(new_text)
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.bold = True  # Make entire line bold
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            elif new_text.strip().startswith("Project Name:") and "Location:" in new_text.strip():
                                # Old format without Duration - still make bold
                                if new_text.strip().startswith("• "):
                                    new_text = new_text.replace("• ", "", 1)
                                run = paragraph.add_run(new_text)
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            else:
                                run = paragraph.add_run(new_text)
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.bold = False
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            # Set 1.5 line spacing
                            paragraph.paragraph_format.line_spacing = 1.5                
                # Remove empty paragraphs from cells
                for p in paragraphs_to_remove:
                    try:
                        p._element.getparent().remove(p._element)
                    except:
                        pass  # Some paragraphs might be required by the table structure
        
        # Third pass: Actually delete the rows (in reverse order to maintain indices)
        for row_idx in sorted(rows_to_delete, reverse=True):
            try:
                row = table.rows[row_idx]
                tbl = table._tbl
                tbl.remove(row._tr)
            except Exception as e:
                st.warning(f"Could not delete row {row_idx}: {str(e)}")
    
    return doc

def safe_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", name).strip() or "output"

# ────────────────────────────────────────────────────────────────
#  Main Application Function
# ────────────────────────────────────────────────────────────────
def main():
    # Check authentication first
    if not check_company_email():
        return
    
    # Check session timeout
    check_session_timeout()
    
    # Display header with user info
    col1, col2 = st.columns([4, 1])
    with col1:
        st.title("📄 CV Company Template Converter")
        st.markdown("Effortlessly reformat candidate CVs into your company's standard template.")
    with col2:
        st.markdown(f"**Logged in as:**  \n{st.session_state.user_email}")
        if st.button("🚪 Logout"):
            # Log logout action
            log_access(st.session_state.user_email, "logout")
            
            # Clear session
            for key in ["authenticated", "user_email", "login_time", "converted_cvs", "conversion_done"]:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()

    # Initialize session state for conversion results
    if 'converted_cvs' not in st.session_state:
        st.session_state.converted_cvs = []
    if 'conversion_done' not in st.session_state:
        st.session_state.conversion_done = False

    # Get API key from secrets
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
    except:
        # Fallback to DEFAULT_API_KEY if secrets not configured
        if DEFAULT_API_KEY:
            api_key = DEFAULT_API_KEY
        else:
            st.error("⚠️ API key not configured. Contact administrator.")
            st.stop()

    # File uploads
    col1, col2 = st.columns(2)
    
    with col1:
        tpl_file = st.file_uploader("Upload Company Template (DOCX)", type=["docx"])
        if tpl_file:
            st.success(f"✅ Template: {tpl_file.name}")
    
    with col2:
        cvs = st.file_uploader("Upload Candidate CV(s)", type=["pdf", "docx", "txt"],
                              accept_multiple_files=True)
        if cvs:
            st.info(f"📁 {len(cvs)} CV(s) uploaded")

    # Process button
    if st.button("🔄 Convert CVs", type="primary", disabled=not(api_key and tpl_file and cvs)):
        # Log conversion attempt
        log_access(st.session_state.user_email, "conversion_started", f"{len(cvs)} CVs")
        
        extractor = CVExtractor(api_key)
        tpl_bytes = tpl_file.getvalue()

        converted = []
        prog = st.progress(0.0)
        status = st.empty()

        for i, cv in enumerate(cvs):
            status.text(f"Processing {cv.name}...")
            
            try:
                # Extract text
                text = extract_text(cv)
                if not text:
                    st.warning(f"⚠️ Could not extract text from {cv.name}")
                    continue
                
                # Extract structured data
                with st.spinner(f"Analyzing {cv.name}..."):
                    data = extractor.extract(text)
                
                # Fill template
                filled = fill_template(Document(BytesIO(tpl_bytes)), data)
                
                # Save to buffer
                buf = BytesIO()
                filled.save(buf)
                buf.seek(0)

                converted.append({
                    "name": data.get("candidate_name", cv.name),
                    "buffer": buf,
                    "data": data
                })
                
                prog.progress((i + 1) / len(cvs))
                
            except Exception as e:
                st.error(f"❌ Error processing {cv.name}: {str(e)}")
                log_access(st.session_state.user_email, "conversion_error", f"{cv.name}: {str(e)}")

        status.empty()
        prog.empty()
        
        if converted:
            st.session_state.converted_cvs = converted
            st.session_state.conversion_done = True
            st.success(f"✅ Successfully converted {len(converted)} CV(s)")
            
            # Log successful conversion
            log_access(st.session_state.user_email, "conversion_success", f"{len(converted)} CVs converted")

    # Display results if conversion is done
    if st.session_state.conversion_done and st.session_state.converted_cvs:
        st.markdown("### Download Converted CVs")
        
        # Option to download all as zip
        if len(st.session_state.converted_cvs) > 1:
            if st.button("📦 Download All as ZIP", type="secondary"):
                import zipfile
                zip_buffer = BytesIO()
                
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for conv in st.session_state.converted_cvs:
                        fname = safe_filename(f"{conv['name']}_Formatted.docx")
                        zip_file.writestr(fname, conv['buffer'].getvalue())
                
                zip_buffer.seek(0)
                st.download_button(
                    "⬇️ Download ZIP Archive",
                    zip_buffer.getvalue(),
                    file_name="converted_cvs.zip",
                    mime="application/zip"
                )
                
                # Log download
                log_access(st.session_state.user_email, "download_zip", f"{len(st.session_state.converted_cvs)} CVs")
        
        # Individual CV downloads
        for idx, conv in enumerate(st.session_state.converted_cvs):
            with st.expander(f"📄 {conv['name']}", expanded=True):
                # Show extracted data summary
                data = conv['data']
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**Extracted Information:**")
                    st.write(f"- Position: {data.get('position', 'N/A')}")
                    st.write(f"- Experience: {data.get('total_experience_years', 'N/A')} years")
                    st.write(f"- Email: {data.get('email', 'N/A')}")
                    st.write(f"- Phone: {data.get('phone', 'N/A')}")
                
                with col2:
                    st.markdown("**Experience Summary:**")
                    actual_experiences = [exp for exp in data.get('experiences', []) 
                                        if exp.get('company') and exp.get('company') != 'N/A']
                    st.write(f"Total experiences: {len(actual_experiences)}")
                    for exp_idx, exp in enumerate(actual_experiences[:3]):
                        st.write(f"{exp_idx+1}. {exp['company']} - {exp.get('role', '')}")
                
                # Download button with unique key
                fname = safe_filename(f"{conv['name']}_Formatted.docx")
                if st.download_button(
                    f"⬇️ Download {fname}",
                    conv['buffer'].getvalue(),
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{idx}"  # Unique key for each button
                ):
                    # Log individual download
                    log_access(st.session_state.user_email, "download_cv", fname)

if __name__ == "__main__":
    main()